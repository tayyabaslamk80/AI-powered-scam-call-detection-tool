"""
Script to convert Thesis_Chapter_5.5_Onwards.txt to Word document (.docx) format
Preserves formatting, headings, and structure
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def is_heading(line):
    """Check if line is a chapter or section heading"""
    if not line.strip():
        return 0
    
    line_clean = line.strip().upper()
    
    # Chapter headings (all caps or title case)
    if line_clean.startswith('CHAPTER '):
        return 1  # Level 1
    if 'CONCLUSIONS' in line_clean and ('&' in line_clean or 'RECOMMENDATIONS' in line_clean):
        return 1
    
    # Numbered sections (1.1, 1.1.1, 2.3.4, etc.)
    if line_clean and line_clean[0].isdigit():
        # Pattern: number.number.number (up to 3 levels)
        parts = line_clean.split('.', 3)
        if len(parts) >= 2:
            first_part = parts[0].strip()
            # Check if first part is a digit
            if first_part.isdigit():
                # Count consecutive numeric parts
                level_count = 1
                for i in range(1, min(len(parts), 4)):
                    # Remove any trailing spaces and check if it's a number
                    part_clean = parts[i].strip().split()[0] if parts[i].strip() else ''
                    if part_clean and part_clean[0].isdigit():
                        level_count += 1
                    else:
                        break
                
                if level_count >= 2 and level_count <= 4:
                    return level_count
    
    # Specific section markers (title case)
    special_headings = [
        'ABSTRACT', 'UNDERTAKING', 'ACKNOWLEDGEMENTS', 'TABLE OF CONTENTS',
        'LIST OF FIGURES', 'REFERENCES', 'ABBREVIATIONS'
    ]
    
    for heading in special_headings:
        if line_clean.startswith(heading):
            return 1
    
    # Check for all-caps headings (likely major sections)
    if line_clean.isupper() and len(line_clean) > 3 and len(line_clean.split()) <= 8:
        return 2
    
    return 0

def is_table_line(line):
    """Check if line appears to be part of a table"""
    # Tables usually have multiple columns separated by tabs or multiple spaces
    if '\t' in line:
        parts = line.split('\t')
        if len(parts) >= 2:
            return True
    # Multiple spaces between words (table format)
    if '  ' in line and line.count('  ') >= 2:
        return True
    return False

def convert_text_to_word(input_file, output_file):
    """Convert text file to Word document"""
    
    # Read the text file
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create Word document
    doc = Document()
    
    # Set default font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    in_table = False
    table_data = []
    prev_was_empty = False
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        line_stripped = line.strip()
        
        # Handle empty lines
        if not line_stripped:
            if in_table and table_data:
                # End current table and add it to document
                add_table_to_doc(doc, table_data)
                table_data = []
                in_table = False
                prev_was_empty = True
            elif not prev_was_empty:
                # Add single paragraph break
                doc.add_paragraph()
                prev_was_empty = True
            i += 1
            continue
        
        prev_was_empty = False
        
        # Check if this is a table line (has tabs or multiple spaced columns)
        if is_table_line(line) and not is_heading(line):
            if not in_table:
                in_table = True
                table_data = []
            
            # Process table line
            if '\t' in line:
                row_data = [cell.strip() for cell in line.split('\t') if cell.strip()]
            else:
                # Try splitting by multiple spaces (2+ spaces)
                import re
                row_data = [cell.strip() for cell in re.split(r' {2,}', line) if cell.strip()]
            
            if row_data and len(row_data) >= 2:
                table_data.append(row_data)
            elif in_table:
                # End table if row doesn't look like table data
                add_table_to_doc(doc, table_data)
                table_data = []
                in_table = False
                # Process this line as regular text
                heading_level = is_heading(line)
                if heading_level > 0:
                    add_heading_to_doc(doc, line_stripped, heading_level)
                else:
                    add_paragraph_to_doc(doc, line_stripped)
            
            i += 1
            continue
        
        # If we were in a table but this isn't a table line, end the table
        if in_table and table_data:
            add_table_to_doc(doc, table_data)
            table_data = []
            in_table = False
        
        # Check for headings
        heading_level = is_heading(line)
        
        if heading_level > 0:
            add_heading_to_doc(doc, line_stripped, heading_level)
        else:
            add_paragraph_to_doc(doc, line_stripped)
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_data:
        add_table_to_doc(doc, table_data)
    
    # Save document
    doc.save(output_file)
    print(f"✓ Successfully converted to: {output_file}")

def add_heading_to_doc(doc, text, level):
    """Add a heading with proper formatting"""
    if level == 1:
        heading = doc.add_heading(text, level=1)
    elif level == 2:
        heading = doc.add_heading(text, level=2)
    elif level == 3:
        heading = doc.add_heading(text, level=3)
    else:
        heading = doc.add_heading(text, level=4)
    
    # Style heading - Times New Roman, Bold
    if heading.runs:
        heading_format = heading.runs[0].font
        heading_format.name = 'Times New Roman'
        heading_format.bold = True
        if level == 1:
            heading_format.size = Pt(16)
        elif level == 2:
            heading_format.size = Pt(14)
        else:
            heading_format.size = Pt(12)

def add_paragraph_to_doc(doc, text):
    """Add a paragraph with proper formatting"""
    para = doc.add_paragraph(text)
    if para.runs:
        para_format = para.runs[0].font
        para_format.name = 'Times New Roman'
        para_format.size = Pt(12)
    return para

def add_table_to_doc(doc, table_data):
    """Add a table to the document"""
    if not table_data:
        return
    
    # Determine number of columns from first row
    max_cols = max(len(row) for row in table_data) if table_data else 1
    
    # Create table
    table = doc.add_table(rows=len(table_data), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for row_idx, row_data in enumerate(table_data):
        for col_idx in range(max_cols):
            cell = table.rows[row_idx].cells[col_idx]
            if col_idx < len(row_data):
                cell.text = row_data[col_idx]
            else:
                cell.text = ""
            
            # Set cell font
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

if __name__ == '__main__':
    input_file = 'Thesis_Chapter_5.5_Onwards.txt'
    output_file = 'Final_Thesis_AI_Powered_Scam_Call_Identification_Tool.docx'
    
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found!")
        exit(1)
    
    print(f"Converting {input_file} to Word format...")
    convert_text_to_word(input_file, output_file)
    print(f"\n✓ Conversion complete! Output: {output_file}")

